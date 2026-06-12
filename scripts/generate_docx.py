# -*- coding: utf-8 -*-
import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Thiết lập màu nền cho ô trong bảng"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập khoảng cách đệm (padding) trong ô"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Thiết lập đường viền mờ cho bảng"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def add_heading_styled(doc, text, level):
    """Thêm tiêu đề với màu sắc và cỡ chữ chuẩn"""
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x60) # Navy Blue
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x3B, 0x7A, 0x57) # Deep Green
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68) # Slate Gray
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(2)
    return heading

def add_paragraph_styled(doc, text="", bold_prefix=None, style='Normal'):
    """Thêm đoạn văn với font chữ Arial mặc định"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.bold = True
        r_bold.font.size = Pt(11)
        r_bold.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

def add_bullet_styled(doc, bold_prefix, text):
    """Thêm danh sách dạng chấm tròn"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    r_bold = p.add_run(bold_prefix)
    r_bold.font.name = 'Arial'
    r_bold.font.bold = True
    r_bold.font.size = Pt(11)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(11)
    return p

def add_diagram_styled(doc, text):
    """Thêm sơ đồ dạng Unicode Art định dạng bằng font chữ Courier New/Consolas"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    
    # Đọc từng dòng để thêm
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Thiết lập màu nền sáng cho đoạn code
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7FAFC"/>')
    pPr.append(shd)
    return p

def create_api_table(doc, apis):
    """Tạo bảng đặc tả danh sách API"""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    headers = ["API / Endpoint", "Xác thực", "Tham số / Payload", "Phản hồi mong đợi"]
    widths = [Inches(2.2), Inches(1.0), Inches(2.0), Inches(1.8)]
    
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F3A60") # Navy background
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        hdr_cells[i].width = widths[i]
        
    for api in apis:
        row = table.add_row()
        row_cells = row.cells
        
        row_cells[0].text = ""
        p0 = row_cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r_verb = p0.add_run(api['verb'] + " ")
        r_verb.font.bold = True
        r_verb.font.name = 'Arial'
        r_verb.font.size = Pt(9.5)
        if api['verb'] in ['POST', 'PUT', 'PATCH']:
            r_verb.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Orange-ish
        elif api['verb'] == 'DELETE':
            r_verb.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) # Red
        else:
            r_verb.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Blue
            
        r_path = p0.add_run(api['path'])
        r_path.font.name = 'Arial'
        r_path.font.size = Pt(9.5)
        
        row_cells[1].text = api['auth']
        p1 = row_cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r_auth = p1.runs[0]
        r_auth.font.name = 'Arial'
        r_auth.font.size = Pt(9.5)
        if "Không" in api['auth']:
            r_auth.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        else:
            r_auth.font.bold = True
            r_auth.font.color.rgb = RGBColor(0x1F, 0x3A, 0x60)
            
        row_cells[2].text = api['payload']
        p2 = row_cells[2].paragraphs[0]
        p2.paragraph_format.space_after = Pt(2)
        p2.runs[0].font.name = 'Arial'
        p2.runs[0].font.size = Pt(9)
        
        row_cells[3].text = api['response']
        p3 = row_cells[3].paragraphs[0]
        p3.paragraph_format.space_after = Pt(2)
        p3.runs[0].font.name = 'Arial'
        p3.runs[0].font.size = Pt(9.5)
        
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.width = widths[i]
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def generate_srs_document(output_path):
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # --- TRANG BÌA ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM (SRS)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x60)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(100)
    run_subtitle = subtitle_p.add_run("HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ SHOPDECORBEE (BEESHOP)\nĐặc Tả Chi Tiết & Sơ Đồ 7 Dịch Vụ Core Backend API")
    run_subtitle.font.name = 'Arial'
    run_subtitle.font.size = Pt(14)
    run_subtitle.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(12)
    run_info = info_p.add_run(
        "Giảng viên hướng dẫn: Nguyễn Văn Chiến\n"
        "Đối tượng tài liệu: Đội ngũ Phát triển, Đội ngũ Kiểm thử (QA/QC)\n"
        "Phiên bản: 1.0.0\n"
        "Ngày lập: 05 tháng 06 năm 2026\n"
        "Trạng thái: Hoàn tất thiết kế - Sẵn sàng kiểm thử"
    )
    run_info.font.name = 'Arial'
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    doc.add_page_break()
    
    # --- PHẦN 1: MỤC TIÊU & TỔNG QUAN ---
    add_heading_styled(doc, "1. Tổng Quan Hệ Thống & Phạm Vi Tài Liệu", level=1)
    add_paragraph_styled(doc, 
        "Tài liệu Đặc tả Yêu cầu Phần mềm (SRS) này được xây dựng cho hệ thống thương mại điện tử ShopdecorBee (chuyên cung cấp sản phẩm trang trí nội thất gia đình). Hệ thống được triển khai trên nền tảng Backend .NET 9 theo kiến trúc Clean Architecture chia làm các tầng API, Application, Domain và Infrastructure. Frontend sử dụng framework Angular. Cơ sở dữ liệu sử dụng Microsoft SQL Server."
    )
    add_paragraph_styled(doc, 
        "Tài liệu tập trung đặc tả chi tiết 7 dịch vụ core backend phục vụ cho quá trình kiểm thử tự động, tích hợp và kiểm thử đơn vị. Toàn bộ đặc tả khớp hoàn toàn với cấu trúc mã nguồn thực tế và kịch bản kiểm thử Postman đã được thiết kế."
    )
    
    # --- PHẦN HỆ THỐNG SƠ ĐỒ ---
    add_heading_styled(doc, "2. Sơ Đồ Hệ Thống & Luồng Nghiệp Vụ (Diagrams)", level=1)
    
    add_heading_styled(doc, "2.1 Sơ đồ Kiến trúc Phân tầng các Dịch vụ (Architecture)", level=2)
    add_paragraph_styled(doc, "Dưới đây là sơ đồ mô tả luồng giao tiếp phân tầng của 7 dịch vụ core từ giao diện phía client cho tới cơ sở dữ liệu SQL Server:")
    
    architecture_art = (
        "┌────────────────────────────────────────────────────────┐\n"
        "│               Client (Angular App / Postman)           │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ (HTTP Requests, Token)\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│               HomeDecorShop.API Controllers            │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ (DI / Abstractions)\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│          HomeDecorShop.Application Services           │\n"
        "│ (AuthSvc, ProductSvc, CategorySvc, CartSvc, OrderSvc) │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ (Entity Mapping via EF)\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│               HomeDecorShop.Infrastructure             │\n"
        "│            (AppDbContext, SQL Repositories)            │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ (EF Core / T-SQL)\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│              SQL Server Database (Docker)              │\n"
        "└────────────────────────────────────────────────────────┘"
    )
    add_diagram_styled(doc, architecture_art)
    
    add_heading_styled(doc, "2.2 Sơ đồ Quan hệ Thực thể Database (Entity Relationship Diagram - ERD)", level=2)
    add_paragraph_styled(doc, "Sơ đồ thể hiện liên kết thực thể dữ liệu trong SQL Server giữa 7 dịch vụ chính:")
    
    erd_art = (
        "  ┌───────────┐             ┌───────────┐             ┌────────────────┐\n"
        "  │   USER    │1 ───────── 1│   CART    │1 ───────── N│   CART_ITEM    │\n"
        "  └─────┬─────┘             └───────────┘             └───────┬────────┘\n"
        "        │1                                                    │N\n"
        "        ├────────────────── N (Sổ địa chỉ Address)            │\n"
        "        │1                                                    │\n"
        "        ├────────────────── 1 (Ví Wallet) ── 1:N ── Transactions\n"
        "        │1                                                    │\n"
        "        └────────────────── N (Đơn đặt hàng Order)            │\n"
        "                              │1                              │\n"
        "                              └───── N (Order_Item) ◄─────────┘\n"
        "                                       │N\n"
        "                                    ┌──┴────────┐\n"
        "                                    │  PRODUCT  │ ◄─── N:1 ─── CATEGORY\n"
        "                                    └───────────┘"
    )
    add_diagram_styled(doc, erd_art)

    add_heading_styled(doc, "2.3 Sơ đồ Phân quyền chức năng theo Tác nhân (Use Case Overview)", level=2)
    add_paragraph_styled(doc, "Phân tách vai trò tương tác giữa Khách vãng lai, Khách đăng nhập và Quản trị viên (Admin):")
    
    usecase_art = (
        "  [Khách Vãng Lai] ──────► Đăng ký, Đăng nhập, Xem sản phẩm, Tìm kiếm & Lọc\n\n"
        "  [Khách Đăng Nhập] ─────► Quản lý hồ sơ, Giỏ hàng, Đặt đơn hàng, Thanh toán,\n"
        "                           Nạp/Rút ví điện tử, Đánh giá sản phẩm\n\n"
        "  [Quản Trị Viên] ───────► Xem ds Users, CRUD Sản phẩm & Danh mục,\n"
        "                           Duyệt đơn hàng, Phê duyệt hoàn tiền giao dịch ví"
    )
    add_diagram_styled(doc, usecase_art)

    add_heading_styled(doc, "2.4 Sơ đồ luồng Thanh toán bằng Ví điện tử (Sequence Flow)", level=2)
    add_paragraph_styled(doc, "Luồng tuần tự khi Khách hàng gửi yêu cầu mua sản phẩm bằng số dư ví điện tử:")
    
    sequence_art = (
        "  Customer               Angular FE           Cart/Order API         Wallet Service\n"
        "     │                       │                      │                      │\n"
        "     │─── 1. Đặt hàng ──────►│                      │                      │\n"
        "     │                       │─── 2. POST /orders ─►│                      │\n"
        "     │                       │                      │── 3. Trừ Stock/Cart ─┤\n"
        "     │                       │◄── 4. OrderView ─────│                      │\n"
        "     │                       │                      │                      │\n"
        "     │                       │─── 5. POST /payments ──────────────────────►│\n"
        "     │                       │      (Method: wallet)│                      │\n"
        "     │                       │                      │◄── 6. Kiểm tra Bal ──│\n"
        "     │                       │                      │                      │\n"
        "     │                       │                      │ [Nếu số dư >= giá]   │\n"
        "     │                       │                      │── 7. Trừ tiền ví ───►│\n"
        "     │                       │                      │   Cập nhật Trạng thái│\n"
        "     │                       │◄── 8. Success ───────┼──────────────────────┤\n"
        "     │◄── 9. Thành công ─────│                      │                      │"
    )
    add_diagram_styled(doc, sequence_art)

    doc.add_page_break()
    
    # --- PHẦN 3: CHI TIẾT 7 DỊCH VỤ ---
    add_heading_styled(doc, "3. Đặc Tả Chi Tiết 7 Dịch Vụ Core", level=1)
    
    # ==================== 3.1 AUTH & USER SERVICE ====================
    add_heading_styled(doc, "3.1. Dịch Vụ Xác Thực & Người Dùng (Auth & User Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Quản lý vòng đời tài khoản người dùng, bao gồm đăng ký tài khoản mới, xác thực thông tin đăng nhập, quản lý hồ sơ cá nhân và phân quyền bảo mật trên toàn bộ hệ thống API.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "User (Người dùng): ", "Bao gồm UserId (PK, int), Email (Unique, string), PasswordHash (string), FullName (string), Phone (string), Address (string), Role (UserRole: Admin=0, Customer=1), CreatedAt (DateTime), IsActive (bool), CurrentToken (string, dùng để lưu session token hiện hành), IsEmailConfirmed (bool), EmailConfirmationToken (string).")
    add_bullet_styled(doc, "Address (Sổ địa chỉ phụ): ", "Id (PK, int), UserId (FK, int), FullName (string), Phone (string), Line1 (Địa chỉ chi tiết, string), Ward (Phường, string), District (Quận, string), City (Thành phố, string), IsDefault (bool).")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    auth_apis = [
        {
            "verb": "POST", "path": "/api/auth/register", "auth": "Không",
            "payload": "RegisterUserInput\n- Email (string, bắt buộc)\n- Password (string, bắt buộc, min 6)\n- FullName (string)\n- Phone (string)",
            "response": "200 OK: Trả về AuthResult { token, role, email }\n400 Bad Request: Dữ liệu không hợp lệ\n409 Conflict: Trùng Email"
        },
        {
            "verb": "POST", "path": "/api/auth/login", "auth": "Không",
            "payload": "LoginInput\n- Email (string, bắt buộc)\n- Password (string, bắt buộc)",
            "response": "200 OK: Trả về AuthResult { token, role, email }\n401 Unauthorized: Sai thông tin tài khoản\n400 Bad Request: Chưa xác thực email"
        },
        {
            "verb": "GET", "path": "/api/auth/confirm-email", "auth": "Không",
            "payload": "Query parameters:\n- token (string, bắt buộc)",
            "response": "200 OK: MessageResponse ('Xac nhan email thanh cong')\n400 Bad Request: Mã xác thực không hợp lệ hoặc hết hạn"
        },
        {
            "verb": "GET", "path": "/api/account/profile", "auth": "Token người dùng",
            "payload": "Headers:\n- Authorization: Bearer <token>\n- Hoặc X-Auth-Token",
            "response": "200 OK: Trả về UserView { id, email, fullName, phone, role, addresses }\n401 Unauthorized: Token hết hạn/không hợp lệ"
        },
        {
            "verb": "PUT", "path": "/api/account/profile", "auth": "Token người dùng",
            "payload": "UpdateProfileInput\n- FullName (string, bắt buộc)\n- Phone (string, bắt buộc)",
            "response": "200 OK: Trả về UserView sau khi cập nhật\n400 Bad Request: Thiếu thông tin"
        },
        {
            "verb": "GET", "path": "/api/users", "auth": "Token Admin",
            "payload": "Headers: Admin Bearer Token",
            "response": "200 OK: Trả về mảng danh sách người dùng\n403 Forbidden: Không phải Admin"
        },
        {
            "verb": "POST", "path": "/api/account/change-password", "auth": "Token người dùng",
            "payload": "ChangePasswordInput\n- OldPassword (string)\n- NewPassword (string)",
            "response": "404 Not Found / 405 Method Not Allowed\nLưu ý: Tính năng chưa được cài đặt thực tế ở backend (dành cho sprint sau)."
        }
    ]
    create_api_table(doc, auth_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Quy chế bảo mật token: ", "Hệ thống sử dụng TokenAuthenticationHandler làm nhiệm vụ đọc token được gửi lên. Một token hợp lệ khi nó khớp với trường CurrentToken của bản ghi User trong DB. Custom Token này giúp đơn giản hóa hệ thống session nhưng yêu cầu kiểm tra DB mỗi khi gọi API.")
    add_bullet_styled(doc, "Đăng ký và Xác nhận: ", "Tài khoản đăng ký mới sẽ tạo ngẫu nhiên một EmailConfirmationToken. Tùy thuộc vào thiết lập hệ thống, tài khoản chưa xác nhận email qua link (/api/auth/confirm-email) sẽ trả về lỗi 400 Bad Request khi cố gắng đăng nhập.")
    
    # ==================== 3.2 PRODUCT SERVICE ====================
    doc.add_page_break()
    add_heading_styled(doc, "3.2. Dịch Vụ Sản Phẩm (Product Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Quản lý danh mục hàng hóa trang trí nội thất bao gồm thông tin chi tiết sản phẩm, quản lý kho số lượng tồn, tìm kiếm, phân trang và đánh giá xếp hạng sản phẩm từ phía khách hàng.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "Product (Sản phẩm): ", "ProductId (PK, int), Sku (Unique, string), ProductName (string), Slug (Unique, string), Price (decimal), OldPrice (decimal, nullable), CategoryId (FK, int), Category (string), Image (string), HoverImage (string), VideoUrl (string, null), Tag (string, null), SoldPercentage (int, null), StockLeft (int), Rating (double), Reviews (int), Brand (string), Color (string), Material (string), Style (string), InStock (bool), IsActive (bool), CreatedAt (DateTime), Description (string, null).")
    add_bullet_styled(doc, "ProductReview (Đánh giá): ", "Id (PK, int), ProductId (FK, int), ReviewerName (string), Rating (int, 1-5), Comment (string), CreatedAt (DateTime).")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    prod_apis = [
        {
            "verb": "GET", "path": "/api/products", "auth": "Không",
            "payload": "Query parameters:\n- q (từ khóa tìm kiếm)\n- category (slug danh mục)\n- brand, style\n- minPrice, maxPrice\n- inStock (bool), onSale (bool)\n- page (mặc định 1), pageSize (mặc định 20)",
            "response": "200 OK: Trả về ProductListResult { items: ProductView[], total: int }\nBao gồm dữ liệu phân trang chuẩn."
        },
        {
            "verb": "GET", "path": "/api/products/{id}", "auth": "Không",
            "payload": "Tham số đường dẫn:\n- id (int, mã sản phẩm)",
            "response": "200 OK: Trả về ProductView chi tiết\n404 NotFound: Không có sản phẩm nào khớp với ID cung cấp"
        },
        {
            "verb": "POST", "path": "/api/products", "auth": "Token Admin",
            "payload": "ProductUpsertInput\n- Sku (string), ProductName (string)\n- Price (decimal), CategoryId (int)\n- StockLeft (int), Brand, Color, v.v.",
            "response": "201 Created: Trả về ProductView vừa tạo\n409 Conflict: Trùng Sku hoặc Slug\n400 BadRequest: Lỗi dữ liệu validation"
        },
        {
            "verb": "PUT", "path": "/api/products/{id}", "auth": "Token Admin",
            "payload": "Tham số đường dẫn: id (int)\nBody: ProductUpsertInput",
            "response": "200 OK: Trả về ProductView sau cập nhật\n404 NotFound: Không có sản phẩm\n409 Conflict: Trùng SKU/Slug của SP khác"
        },
        {
            "verb": "DELETE", "path": "/api/products/{id}", "auth": "Token Admin",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "204 No Content: Xóa thành công\n404 NotFound: Không tìm thấy sản phẩm"
        },
        {
            "verb": "GET", "path": "/api/products/{id}/reviews", "auth": "Không",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "200 OK: Trả về danh sách ProductReviewView[]"
        },
        {
            "verb": "POST", "path": "/api/products/{id}/reviews", "auth": "Token người dùng",
            "payload": "ProductReviewCreateInput\n- ProductId (int)\n- ReviewerName (string)\n- Rating (int, từ 1 đến 5)\n- Comment (string)",
            "response": "200 OK: Trả về ProductReviewView\n400 BadRequest: Sai ProductId trong body so với path"
        }
    ]
    create_api_table(doc, prod_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Duy nhất mã SKU và Slug: ", "Hệ thống tự động chuyển đổi ProductName thành Slug dạng không dấu kết nối bằng dấu gạch ngang (ví dụ: 'Ke sach bee' -> 'ke-sach-bee'). Sku và Slug phải là duy nhất trên toàn hệ thống sản phẩm. Nếu vi phạm, trả về lỗi 409 Conflict.")
    add_bullet_styled(doc, "Ràng buộc số lượng tồn kho (Stock): ", "Nếu StockLeft <= 0, InStock sẽ tự động chuyển thành false. Khi tạo đơn hàng mới, số lượng sản phẩm mua sẽ trừ trực tiếp vào StockLeft. Nếu vượt quá số lượng còn lại, hệ thống sẽ báo lỗi hết hàng (Conflict).")
    
    # ==================== 3.3 CATEGORY SERVICE ====================
    doc.add_page_break()
    add_heading_styled(doc, "3.3. Dịch Vụ Danh Mục (Category Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Phân loại hàng hóa theo các tầng danh mục và nhóm danh mục sản phẩm nhằm tối ưu hóa trải nghiệm tìm kiếm và bộ lọc phía khách hàng.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "Category (Danh mục sản phẩm): ", "Id (PK, int), Name (string), Slug (Unique, string), IsActive (bool), GroupId (FK, int).")
    add_bullet_styled(doc, "CategoryGroup (Nhóm danh mục): ", "Id (PK, int), Name (string), Slug (Unique, string), IsActive (bool).")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    cat_apis = [
        {
            "verb": "GET", "path": "/api/categories", "auth": "Không",
            "payload": "Không cần tham số",
            "response": "200 OK: Trả về mảng CategoryView[]\nThông tin CategoryGroup được lồng trực tiếp trong mỗi Category qua navigation property 'GroupNavigation'."
        },
        {
            "verb": "GET", "path": "/api/categories/{id}", "auth": "Không",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "200 OK: Trả về CategoryView cụ thể\n404 NotFound: Không tìm thấy danh mục"
        },
        {
            "verb": "POST", "path": "/api/categories", "auth": "Token Admin",
            "payload": "CategoryUpsertInput\n- Name (string)\n- Slug (string)\n- GroupId (int)\n- IsActive (bool)",
            "response": "201 Created: Trả về danh mục vừa tạo thành công\n409 Conflict: Trùng tên hoặc Slug danh mục"
        },
        {
            "verb": "PUT", "path": "/api/categories/{id}", "auth": "Token Admin",
            "payload": "Tham số đường dẫn: id (int)\nBody: CategoryUpsertInput",
            "response": "200 OK: Trả về danh mục đã sửa đổi\n404 NotFound: Danh mục không tồn tại\n409 Conflict: Trùng tên/Slug hoặc hành vi vô hiệu hóa vi phạm ràng buộc sản phẩm"
        },
        {
            "verb": "DELETE", "path": "/api/categories/{id}", "auth": "Token Admin",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "204 No Content: Xóa thành công\n409 Conflict: Danh mục hiện đang liên kết với các sản phẩm trong hệ thống\n404 NotFound: Không tìm thấy danh mục"
        }
    ]
    create_api_table(doc, cat_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Nhóm danh mục lồng ghép: ", "Trong hệ thống hiện tại, không có endpoint riêng biệt để tương tác độc lập với CategoryGroup (như cập nhật, xóa nhóm danh mục qua HTTP). Thông tin nhóm danh mục được quản lý và hiển thị lồng trong kết quả của danh mục sản phẩm.")
    add_bullet_styled(doc, "Ràng buộc xóa danh mục (Deletion safety): ", "Khi xóa một danh mục thông qua DELETE /api/categories/{id}, hệ thống sẽ kiểm tra xem có sản phẩm nào thuộc danh mục này hay không. Nếu có, API trả về mã lỗi 409 Conflict kèm mã lỗi nội bộ CategoryHasProducts để đảm bảo toàn vẹn dữ liệu cơ sở dữ liệu.")
    
    # ==================== 3.4 CART SERVICE ====================
    doc.add_page_break()
    add_heading_styled(doc, "3.4. Dịch Vụ Giỏ Hàng (Cart Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Lưu trữ tạm thời các sản phẩm người dùng chọn mua trước khi thanh toán. Dịch vụ hỗ trợ thêm, sửa đổi số lượng hàng hóa và tự động tính toán giá trị tạm tính.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "Cart (Giỏ hàng tổng): ", "Id (PK, int), UserId (FK, int, liên kết 1-1 với User), CreatedAt (DateTime), UpdatedAt (DateTime).")
    add_bullet_styled(doc, "CartItem (Chi tiết mặt hàng trong giỏ): ", "Id (PK, int), CartId (FK, int), ProductId (FK, int), Quantity (int), UnitPrice (decimal), CreatedAt (DateTime), UpdatedAt (DateTime).")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    cart_apis = [
        {
            "verb": "GET", "path": "/api/cart", "auth": "Token người dùng",
            "payload": "Không cần tham số body",
            "response": "200 OK: Trả về CartView { id, items: CartItemView[] }\nHệ thống tự khởi tạo giỏ hàng trống nếu là lần truy cập đầu tiên."
        },
        {
            "verb": "POST", "path": "/api/cart/items", "auth": "Token người dùng",
            "payload": "AddCartItemInput\n- ProductId (int, bắt buộc)\n- Quantity (int, bắt buộc, > 0)",
            "response": "200 OK: Trả về CartView mới\n404 NotFound: Sản phẩm không tồn tại\n409 Conflict: Số lượng vượt quá giới hạn tồn kho sản phẩm"
        },
        {
            "verb": "PUT", "path": "/api/cart/items/{itemId}", "auth": "Token người dùng",
            "payload": "Tham số đường dẫn: itemId (int)\nBody: UpdateCartItemQuantityInput\n- Quantity (int, bắt buộc, > 0)",
            "response": "200 OK: Trả về CartView sau khi đổi số lượng\n404 NotFound: Không tìm thấy mặt hàng trong giỏ\n409 Conflict: Vượt quá số hàng còn lại trong kho"
        },
        {
            "verb": "DELETE", "path": "/api/cart/items/{itemId}", "auth": "Token người dùng",
            "payload": "Tham số đường dẫn: itemId (int)",
            "response": "204 No Content: Xóa mặt hàng khỏi giỏ thành công\n404 NotFound: Không tìm thấy mặt hàng để xóa"
        },
        {
            "verb": "DELETE", "path": "/api/cart/items", "auth": "Token người dùng",
            "payload": "Không cần tham số",
            "response": "204 No Content: Xóa sạch toàn bộ sản phẩm trong giỏ"
        }
    ]
    create_api_table(doc, cart_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Tính gộp số lượng: ", "Khi thêm một sản phẩm đã có sẵn trong giỏ thông qua POST /api/cart/items, hệ thống không tạo dòng mới mà tự động tăng số lượng hiện tại (Quantity = Quantity + Input.Quantity).")
    add_bullet_styled(doc, "Kiểm tra giới hạn tồn kho: ", "Tất cả các tác vụ thêm hoặc cập nhật số lượng đều phải đối chiếu với số lượng sản phẩm còn lại trong kho (StockLeft). Nếu tổng số lượng yêu cầu lớn hơn StockLeft, hệ thống sẽ từ chối và trả về lỗi 409 Conflict.")
    
    # ==================== 3.5 ORDER SERVICE ====================
    doc.add_page_break()
    add_heading_styled(doc, "3.5. Dịch Vụ Đơn Hàng (Order Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Quản lý toàn bộ quy trình mua hàng từ khi chuyển đổi giỏ hàng thành đơn hàng, theo dõi các trạng thái xử lý, cho phép hủy đơn và xử lý khiếu nại hoàn tiền.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "Order (Đơn hàng): ", "Id (PK, int), UserId (FK, int), OrderNumber (Unique, string), Status (OrderStatus: PendingPayment=0, Processing=1, Cancelled=2, Completed=3, RefundRequested=4, Refunded=5), PaymentStatus (PaymentStatus: Pending=0, Paid=1, Failed=2, Cancelled=3, Refunded=4), Subtotal (decimal), ShippingFee (decimal), TotalAmount (Subtotal + ShippingFee, decimal), FullName (string), Phone (string), Line1 (string), Ward (string), District (string), City (string), Notes (string), CreatedAt, UpdatedAt.")
    add_bullet_styled(doc, "OrderItem (Chi tiết đơn hàng): ", "Id (PK, int), OrderId (FK, int), ProductId (FK, int), Quantity (int), Price (decimal - Giá chụp lại lúc mua), ProductName (string - Tên sản phẩm chụp lại lúc mua).")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    order_apis = [
        {
            "verb": "POST", "path": "/api/orders", "auth": "Token người dùng",
            "payload": "PlaceOrderInput\n- FullName (string), Phone (string)\n- Line1 (string), Ward, District, City\n- ShippingFee (decimal)\n- Notes (string, tùy chọn)",
            "response": "201 Created: Trả về OrderView chi tiết đơn hàng vừa tạo\n400 BadRequest: Giỏ hàng trống hoặc thiếu thông tin giao hàng\n409 Conflict: Một hoặc nhiều sản phẩm trong giỏ đã hết hàng"
        },
        {
            "verb": "GET", "path": "/api/orders", "auth": "Token người dùng",
            "payload": "Không cần tham số",
            "response": "200 OK: Trả về mảng danh sách đơn hàng OrderView[] của bản thân"
        },
        {
            "verb": "GET", "path": "/api/orders/{id}", "auth": "Token người dùng",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "200 OK: Trả về OrderView cụ thể\n404 NotFound: Đơn hàng không tồn tại hoặc không thuộc người dùng này"
        },
        {
            "verb": "POST", "path": "/api/orders/{id}/cancel", "auth": "Token người dùng",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "200 OK: Hủy thành công, trả về OrderView ở trạng thái Cancelled\n409 Conflict: Đơn hàng ở trạng thái không cho phép hủy (ví dụ: đã giao, đã thanh toán)"
        },
        {
            "verb": "POST", "path": "/api/orders/{id}/request-refund", "auth": "Token người dùng",
            "payload": "Body: RequestRefundInput\n- Reason (string, lý do hoàn tiền)",
            "response": "200 OK: Đổi trạng thái sang RefundRequested\n409 Conflict: Đơn hàng chưa thanh toán hoặc trạng thái không hợp lý"
        },
        {
            "verb": "GET", "path": "/api/admin/orders", "auth": "Token Admin",
            "payload": "Không cần tham số",
            "response": "200 OK: Trả về toàn bộ đơn hàng của hệ thống OrderView[]"
        },
        {
            "verb": "PATCH", "path": "/api/admin/orders/{id}/status", "auth": "Token Admin",
            "payload": "Query parameters:\n- status (string: 'processing', 'completed', 'cancelled')",
            "response": "200 OK: Cập nhật thành công, trả về OrderView mới\n400 BadRequest: Tên trạng thái không hợp lệ\n404 NotFound: Không tìm thấy đơn hàng"
        },
        {
            "verb": "POST", "path": "/api/admin/orders/{id}/process-refund", "auth": "Token Admin",
            "payload": "Query parameters:\n- approve (bool: true/false)",
            "response": "200 OK: Đơn hàng được cập nhật trạng thái hoàn tiền (Refunded nếu approve=true)\n409 Conflict: Đơn hàng chưa yêu cầu hoàn tiền hoặc sai luồng nghiệp vụ"
        }
    ]
    create_api_table(doc, order_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Sao chụp giá và tên sản phẩm: ", "Khi tạo đơn hàng, hệ thống chụp lại (snapshot) giá hiện hành (`Price`) và tên sản phẩm (`ProductName`) lưu vào bảng `OrderItem`. Điều này giúp đảm bảo lịch sử doanh thu không bị ảnh hưởng khi quản trị viên cập nhật thông tin sản phẩm sau này.")
    add_bullet_styled(doc, "Quy trình trừ kho và hoàn kho: ", "Khi một đơn hàng được tạo, số lượng mua sẽ trừ vào tồn kho sản phẩm. Nếu đơn hàng bị hủy (`Cancelled`), hệ thống tự động cộng trả lại số lượng tương ứng cho kho hàng sản phẩm.")
    add_bullet_styled(doc, "Luồng trạng thái đơn hàng (State machine): ", "Đơn hàng mới tạo có trạng thái mặc định là PendingPayment (0) và PaymentStatus = Pending (0). Khi thanh toán thành công, đơn hàng tự động đổi sang Processing (1) và PaymentStatus = Paid (1). Admin có quyền chuyển trạng thái đơn hàng sang Completed (3) hoặc Cancelled (2).")
    
    # ==================== 3.6 PAYMENT SERVICE ====================
    doc.add_page_break()
    add_heading_styled(doc, "3.6. Dịch Vụ Thanh Toán (Payment Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Xử lý giao dịch thanh toán cho đơn hàng. Hỗ trợ phương thức thu tiền khi giao hàng (COD), thanh toán qua cổng VNPay bằng cách tạo URL giao dịch và xử lý phản hồi từ VNPay.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "Payment (Giao dịch thanh toán): ", "Id (PK, int), OrderId (FK, int), Method (string: 'cod', 'wallet', 'vnpay'), Status (PaymentStatus: Pending=0, Paid=1, Failed=2, Cancelled=3, Refunded=4), Amount (decimal), TransactionCode (Unique, string), PaidAt (DateTime, nullable), CreatedAt, UpdatedAt.")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    pay_apis = [
        {
            "verb": "POST", "path": "/api/payments", "auth": "Token người dùng",
            "payload": "PaymentProcessInput\n- OrderId (int, bắt buộc)\n- Method (string: 'cod' hoặc 'wallet')",
            "response": "201 Created: Trả về PaymentView thông tin thanh toán đã tạo\n409 Conflict: Đơn hàng đã được thanh toán hoặc đã bị hủy"
        },
        {
            "verb": "GET", "path": "/api/payments", "auth": "Token người dùng",
            "payload": "Không cần tham số",
            "response": "200 OK: Trả về danh sách PaymentView[] của bản thân"
        },
        {
            "verb": "GET", "path": "/api/payments/{id}", "auth": "Token người dùng",
            "payload": "Tham số đường dẫn: id (int)",
            "response": "200 OK: Trả về PaymentView chi tiết\n404 NotFound: Không tìm thấy giao dịch thanh toán"
        },
        {
            "verb": "GET", "path": "/api/payments/order/{orderId}", "auth": "Token người dùng",
            "payload": "Tham số đường dẫn: orderId (int)",
            "response": "200 OK: Trả về danh sách PaymentView[] của đơn hàng"
        },
        {
            "verb": "POST", "path": "/api/payments/vnpay/url", "auth": "Token người dùng",
            "payload": "VnPayCreateUrlInput\n- OrderId (int, bắt buộc)",
            "response": "201 Created: Trả về VnPayCreateUrlApiResult { paymentId, paymentUrl, ... }\nNgười dùng sẽ được chuyển hướng sang URL này để tiến hành thanh toán."
        },
        {
            "verb": "GET", "path": "/api/payments/vnpay/return", "auth": "Không (VNPay Redirect)",
            "payload": "Query parameters từ VNPay trả về:\n- vnp_TxnRef, vnp_ResponseCode, vnp_SecureHash, vnp_Amount, vnp_TransactionNo, v.v.",
            "response": "302 Found (Redirect): Chuyển hướng về trang Frontend của khách hàng kèm các tham số kết quả (success/failed)"
        },
        {
            "verb": "GET", "path": "/api/payments/vnpay/ipn", "auth": "Không (VNPay Server to Server)",
            "payload": "Query parameters tương tự return callback",
            "response": "200 OK: Trả về JSON VnPayIpnResponse { RspCode, Message } tương thích với chuẩn VNPay (00: Success, 97: Invalid Signature...)"
        }
    ]
    create_api_table(doc, pay_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Xử lý bảo mật VNPay (HMAC-SHA512): ", "Mọi phản hồi từ VNPay (Return và IPN) đều phải xác thực tính toàn vẹn thông qua mã băm bảo mật. Hệ thống sắp xếp các tham số bắt đầu bằng 'vnp_' theo thứ tự bảng chữ cái, kết nối bằng ký tự '&' và tính toán mã HMAC-SHA512 với HashSecret của cấu hình. Nếu không hợp lệ, trả về mã lỗi 97 (Sai chữ ký) cho VNPay.")
    add_bullet_styled(doc, "Cập nhật liên đới trạng thái đơn hàng: ", "Khi một giao dịch thanh toán thành công (PaymentStatus chuyển sang Paid), hệ thống tự động cập nhật PaymentStatus của đơn hàng tương ứng thành Paid (1) và OrderStatus của đơn hàng thành Processing (1) để bắt đầu giao hàng.")
    
    # ==================== 3.7 WALLET SERVICE ====================
    doc.add_page_break()
    add_heading_styled(doc, "3.7. Dịch Vụ Ví Điện Tử (Wallet Service)", level=2)
    add_paragraph_styled(doc, "Mục tiêu: ", bold_prefix="Cung cấp ví tiền ảo tích hợp trong tài khoản để người dùng nạp tiền, rút tiền, thanh toán đơn hàng bằng ví thay thế cho tiền mặt, và tự động nhận tiền hoàn trả khi hủy đơn hàng.")
    
    add_heading_styled(doc, "Thành phần dữ liệu (Database Entities)", level=3)
    add_bullet_styled(doc, "Wallet (Ví điện tử): ", "Id (PK, int), UserId (FK, int, 1-1 với User), Balance (decimal - Số dư ví), UpdatedAt.")
    add_bullet_styled(doc, "WalletTransaction (Lịch sử giao dịch ví): ", "Id (PK, int), WalletId (FK, int), Amount (decimal), Type (WalletTransactionType: Deposit=0, Withdraw=1, Payment=2, Refund=3), Status (WalletTransactionStatus: Pending=0, Success=1, Failed=2), Reference (string - Mã tham chiếu như mã đơn hàng hoặc mã VNPay), Description (string), CreatedAt.")
    
    add_heading_styled(doc, "Danh sách API Đặc Tả", level=3)
    wallet_apis = [
        {
            "verb": "GET", "path": "/api/wallet", "auth": "Token người dùng",
            "payload": "Không cần tham số",
            "response": "200 OK: Trả về thông tin ví WalletView { id, userId, balance, updatedAt }\nTự động tạo ví mới với số dư ban đầu bằng 0 nếu chưa tồn tại."
        },
        {
            "verb": "GET", "path": "/api/wallet/transactions", "auth": "Token người dùng",
            "payload": "Không cần tham số",
            "response": "200 OK: Trả về danh sách WalletTransactionView[]"
        },
        {
            "verb": "POST", "path": "/api/wallet/deposit", "auth": "Token người dùng",
            "payload": "WalletDepositInput\n- Amount (decimal, bắt buộc, tối thiểu 1,000 VND)",
            "response": "200 OK: Trả về thông tin ví sau nạp tiền (Tính năng nạp tiền trực tiếp dùng để test)"
        },
        {
            "verb": "POST", "path": "/api/wallet/deposit/vnpay/url", "auth": "Token người dùng",
            "payload": "WalletDepositInput\n- Amount (decimal, bắt buộc)",
            "response": "200 OK: Trả về WalletDepositUrlResult { reference, paymentUrl } để người dùng thực hiện nạp tiền qua cổng VNPay"
        },
        {
            "verb": "POST", "path": "/api/wallet/withdraw", "auth": "Token người dùng",
            "payload": "WalletWithdrawInput\n- Amount (decimal, bắt buộc)",
            "response": "200 OK: Trả về ví đã rút tiền\n409 Conflict: Số dư ví không đủ để thực hiện rút"
        },
        {
            "verb": "POST", "path": "/api/wallet/pay", "auth": "Token người dùng",
            "payload": "WalletPayOrderInput\n- OrderId (int, bắt buộc)",
            "response": "200 OK: Trả về WalletView đã thanh toán\n404 NotFound: Đơn hàng không tồn tại\n409 Conflict: Số dư ví không đủ để thanh toán giá trị đơn hàng"
        }
    ]
    create_api_table(doc, wallet_apis)
    
    add_heading_styled(doc, "Quy tắc Nghiệp vụ (Business Rules)", level=3)
    add_bullet_styled(doc, "Ràng buộc số dư không âm: ", "Khi thực hiện hành vi rút tiền (Withdraw) hoặc thanh toán đơn hàng bằng ví (PayOrder), hệ thống kiểm tra số dư ví hiện hành (`Balance`). Nếu Balance < Amount, giao dịch sẽ không được thực hiện, bản ghi giao dịch ví được ghi nhận trạng thái Failed (2) và API trả lời mã lỗi 409 Conflict.")
    add_bullet_styled(doc, "Tự động hoàn tiền vào ví (Refund): ", "Khi khách hàng khiếu nại hoàn tiền và được Admin phê duyệt thông qua POST /api/admin/orders/{id}/process-refund?approve=true, nếu đơn hàng đó trước đây được thanh toán bằng Ví điện tử (`wallet`), hệ thống sẽ tự động tạo một giao dịch ví mới loại `Refund` (3) với trạng thái `Success` (1) và cộng trực tiếp số tiền hoàn trả vào số dư ví của khách hàng.")
    
    # --- PHẦN 4: KẾ HOẠCH KIỂM THỬ POSTMAN ---
    doc.add_page_break()
    add_heading_styled(doc, "4. Ánh Xạ Với Kịch Bản Kiểm Thử Postman (Tuần 1)", level=1)
    add_paragraph_styled(doc, 
        "Dưới đây là bảng ánh xạ giữa các yêu cầu đặc tả (SRS) ở trên với các ca kiểm thử cụ thể trong bộ sưu tập Postman 'BeeShop_Week1.postman_collection.json' đã thực hiện:"
    )
    
    # Tạo bảng ánh xạ test case
    t_map = doc.add_table(rows=1, cols=4)
    t_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_map)
    
    map_headers = ["Mã Ca Kiểm Thử", "Dịch Vụ", "Nội Dung Xác Minh", "Yêu Cầu Kết Quả API"]
    map_widths = [Inches(1.2), Inches(1.8), Inches(2.2), Inches(1.8)]
    
    hdr_cells = t_map.rows[0].cells
    for i, title in enumerate(map_headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "3B7A57") # Deep Green background
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        hdr_cells[i].width = map_widths[i]
        
    test_cases_map = [
        {"id": "AUTH-01", "service": "Auth & User", "desc": "Seed toàn bộ dữ liệu mẫu", "expected": "POST /api/Maintenance/seed/all -> 200 OK"},
        {"id": "AUTH-02", "service": "Auth & User", "desc": "Đăng ký tài khoản khách hàng mới", "expected": "POST /api/auth/register -> 200 OK"},
        {"id": "AUTH-03", "service": "Auth & User", "desc": "Đăng nhập tài khoản khách hàng", "expected": "POST /api/auth/login -> 200 OK"},
        {"id": "AUTH-04", "service": "Auth & User", "desc": "Xem hồ sơ cá nhân hiện tại", "expected": "GET /api/account/profile -> 200 OK"},
        {"id": "AUTH-05", "service": "Auth & User", "desc": "Cập nhật tên và số điện thoại", "expected": "PUT /api/account/profile -> 200 OK"},
        {"id": "AUTH-06", "service": "Auth & User", "desc": "Đăng nhập với vai trò Admin", "expected": "POST /api/auth/login (admin) -> 200 OK"},
        {"id": "AUTH-07", "service": "Auth & User", "desc": "Xem danh sách người dùng", "expected": "GET /api/users (Admin-only) -> 200 OK"},
        {"id": "AUTH-08", "service": "Auth & User", "desc": "Yêu cầu đổi mật khẩu (chưa làm)", "expected": "POST /api/account/change-password -> 404/405"},
        
        {"id": "PROD-01", "service": "Product", "desc": "Lấy danh sách sản phẩm phân trang", "expected": "GET /api/products -> 200 OK"},
        {"id": "PROD-02", "service": "Product", "desc": "Lấy chi tiết một sản phẩm", "expected": "GET /api/products/{id} -> 200 OK"},
        {"id": "PROD-03", "service": "Product", "desc": "Tìm kiếm sản phẩm theo từ khóa", "expected": "GET /api/products?q=bee -> 200 OK"},
        {"id": "PROD-04", "service": "Product", "desc": "Bộ lọc sản phẩm theo danh mục", "expected": "GET /api/products?category={slug} -> 200 OK"},
        
        {"id": "CAT-01", "service": "Category", "desc": "Lấy toàn bộ danh sách danh mục", "expected": "GET /api/categories -> 200 OK"},
        {"id": "CAT-02", "service": "Category", "desc": "Xem chi tiết một danh mục", "expected": "GET /api/categories/{id} -> 200 OK"},
        
        {"id": "CART-01", "service": "Cart", "desc": "Xem thông tin giỏ hàng hiện tại", "expected": "GET /api/cart -> 200 OK"},
        {"id": "CART-02", "service": "Cart", "desc": "Thêm mặt hàng mới vào giỏ hàng", "expected": "POST /api/cart/items -> 200 OK"},
        {"id": "CART-03", "service": "Cart", "desc": "Cập nhật số lượng của mặt hàng", "expected": "PUT /api/cart/items/{id} -> 200 OK"},
        {"id": "CART-04", "service": "Cart", "desc": "Xóa một mặt hàng khỏi giỏ", "expected": "DELETE /api/cart/items/{id} -> 204 No Content"},
        {"id": "CART-05", "service": "Cart", "desc": "Xóa sạch toàn bộ giỏ hàng", "expected": "DELETE /api/cart/items -> 204 No Content"},
        
        {"id": "ORD-02", "service": "Order", "desc": "Đặt đơn hàng từ giỏ hàng", "expected": "POST /api/orders -> 201 Created"},
        {"id": "ORD-03", "service": "Order", "desc": "Xem danh sách đơn hàng đã mua", "expected": "GET /api/orders -> 200 OK"},
        {"id": "ORD-05", "service": "Order", "desc": "Admin xem toàn bộ đơn hàng", "expected": "GET /api/admin/orders -> 200 OK"},
        {"id": "ORD-06", "service": "Order", "desc": "Admin cập nhật trạng thái đơn", "expected": "PATCH /api/admin/orders/{id}/status -> 200 OK"},
        {"id": "ORD-07", "service": "Order", "desc": "Khách hàng tự hủy đơn hàng chưa trả tiền", "expected": "POST /api/orders/{id}/cancel -> 200 OK"},
        
        {"id": "PAY-01", "service": "Payment", "desc": "Thực hiện thanh toán đơn hàng (COD)", "expected": "POST /api/payments -> 201 Created"},
        {"id": "PAY-04", "service": "Payment", "desc": "Xem lịch sử thanh toán theo mã đơn", "expected": "GET /api/payments/order/{orderId} -> 200 OK"},
        
        {"id": "WAL-01", "service": "Wallet", "desc": "Kiểm tra ví điện tử và số dư", "expected": "GET /api/wallet -> 200 OK"},
        {"id": "WAL-02", "service": "Wallet", "desc": "Nạp tiền thử nghiệm trực tiếp", "expected": "POST /api/wallet/deposit -> 200 OK"},
        {"id": "WAL-03", "service": "Wallet", "desc": "Yêu cầu rút tiền khỏi ví", "expected": "POST /api/wallet/withdraw -> 200 OK"},
        {"id": "WAL-04", "service": "Wallet", "desc": "Xem lịch sử giao dịch ví", "expected": "GET /api/wallet/transactions -> 200 OK"}
    ]
    
    for tc in test_cases_map:
        row = t_map.add_row()
        row_cells = row.cells
        
        row_cells[0].text = tc['id']
        row_cells[1].text = tc['service']
        row_cells[2].text = tc['desc']
        row_cells[3].text = tc['expected']
        
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.width = map_widths[i]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            if i == 0:
                run.font.bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_paragraph_styled(doc, 
        "Tài liệu kết thúc tại đây. Kính trình giảng viên và đội ngũ dự án phê duyệt."
    )
    
    doc.save(output_path)
    print(f"Successfully generated DOCX with Diagrams at: {output_path}")

if __name__ == "__main__":
    output_file = r"c:\Users\PC\OneDrive\Documents\ShopdecorBee-main\ShopdecorBee_SRS_7_Services.docx"
    generate_srs_document(output_file)
